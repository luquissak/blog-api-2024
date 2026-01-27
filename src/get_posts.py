import sys
import os
from dotenv import load_dotenv
from googleapiclient import sample_tools
from docx import Document
import html2text
from datetime import datetime


load_dotenv()
BLOG_ID = os.getenv('BLOG_ID')
print("BLOG_ID", BLOG_ID)


def main(argv):
    # Authenticate and construct service.
    service, flags = sample_tools.init(
        argv,
        "blogger",
        "v3",
        __doc__,
        __file__,
        scope="https://www.googleapis.com/auth/blogger",
    )
    print("logged!")

    users = service.users()
    thisuser = users.get(userId="self").execute()
    print("This user's display name is: %s" % thisuser["displayName"])
    blogs = service.blogs()
    thisusersblogs = blogs.listByUser(userId="self").execute()
    for blog in thisusersblogs["items"]:
        print("The blog named '%s' is at: %s" % (blog["name"], blog["url"]))

    posts = service.posts()

    document = Document()
    document_2013 = Document()
    document_2014 = Document()
    document_2015 = Document()
    document_2016 = Document()
    document_2017 = Document()
    document_2018 = Document()
    document_2019 = Document()
    document_2020 = Document()
    document_2021 = Document()
    document_2022 = Document()
    document_2023 = Document()
    document_2024 = Document()
    document_2025 = Document()

    i = 0
    for blog in thisusersblogs["items"]:
        print("The posts for %s:" % blog["name"])
        request = posts.list(blogId=blog["id"])
        while request != None:
            posts_doc = request.execute()
            if "items" in posts_doc and not (posts_doc["items"] is None):
                for post in posts_doc["items"]:
                    try:
                        post_day = datetime.strptime(post["published"], '%Y-%m-%dT%H:%M:%S-03:00').date().strftime('%d/%m/%Y')
                        year = datetime.strptime(post["published"], '%Y-%m-%dT%H:%M:%S-03:00').date().strftime('%Y')
                    except:
                        post_day = datetime.strptime(post["published"], '%Y-%m-%dT%H:%M:%S-02:00').date().strftime('%d/%m/%Y')
                        year = datetime.strptime(post["published"], '%Y-%m-%dT%H:%M:%S-02:00').date().strftime('%Y')
                    print("%s (%s)..." % (post["title"], post_day))
                    print("%s ..." % (year))
                    if year == "2013":
                        document_2013.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2013.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2013.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    elif year == "2014":
                        document_2014.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2014.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2014.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    elif year == "2015":
                        document_2015.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2015.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2015.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    elif year == "2016":
                        document_2016.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2016.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2016.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    elif year == "2017":
                        document_2017.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2017.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2017.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    elif year == "2018":
                        document_2018.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2018.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2018.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    elif year == "2019":
                        document_2019.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2019.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2019.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    elif year == "2020":
                        document_2020.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2020.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2020.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    elif year == "2021":
                        document_2021.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2021.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2021.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    elif year == "2022":
                        document_2022.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2022.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2022.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    elif year == "2023":
                        document_2023.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2023.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2023.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    elif year == "2024":
                        document_2024.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2024.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2024.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    elif year == "2025":
                        document_2025.add_heading(
                        post["title"]+" - "+post_day+"\n", level=1)
                        document_2025.add_paragraph(html2text.html2text(post["content"])+"\n")
                        document_2025.save('docao3/Reflexoes_'+str(year)+'_posts_content.docx')
                    else:
                        print("DEU RUIM")
#                    continue
                    #if i > 3: exit()
                    document.add_heading(
                       post["title"]+" - "+post_day+"\n", level=1)
                    document.add_paragraph(html2text.html2text(post["content"])+"\n")
                    document.save('docao3/Reflexoes_posts_content.docx')
                    i = i + 1
            request = posts.list_next(request, posts_doc)


if __name__ == "__main__":
    main(sys.argv)
